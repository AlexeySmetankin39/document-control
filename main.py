from docx import Document

doc = Document("sample.docx") #загрузка документа
pattern_id = [] #список для хранения позиций и индексов параграфов
substring = "***"   #подстрока для поиска

for i , paragraphs in enumerate(doc.paragraphs): #пройти по всем параграфам документа
    text = paragraphs.text.strip() #получить текст параграфа
    if text: #проверить, что текст не пустой
        if substring in text: #проверить, что подстрока присутствует в тексте
            position = text.find(substring) #найти позицию подстроки
            pattern_id.append({i, position}) #добавить позицию и индекс параграфа в список


graph_name = ["название работы: ", "Статья ДДС: ", "Дата:"]
text_array = []
for i in range(len(pattern_id)):
    text_array.append(input("Введите " + graph_name[i])) #ввод текста в массив

for i in range(len(pattern_id)):
    paragraph = doc.paragraphs[list(pattern_id[i])[1]] #получить параграф по индексу
    text = paragraph.text
    new_text = text.replace(substring, "") #заменить подстроку на пустую строку
    paragraph.text = new_text


for i in range(len(pattern_id)):
    paragraph = doc.paragraphs[list(pattern_id[i])[1]] #получить параграф по индексу
    text = paragraph.text #получить текст параграфа
    new_text = text.replace(substring, "") #заменить подстроку на пустую строку
    text_before = paragraph.text[:list(pattern_id[i])[0]]
    text_after = paragraph.text[list(pattern_id[i])[0]:]
    paragraph.text = text_before + text_array[i] + text_after
doc.save("измененный_файл.docx")

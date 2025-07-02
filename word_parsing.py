from docx import Document
import os
from datetime import datetime
from io import BytesIO

WORK_NAME_SUBSTRING = '__WORK__NAME__'
DDS_SUBSTRING = '__DDS__'
DATE_SUBSTRING = '__DATE__'

def searc_text_position(doc):
    """Находит позиции ключевых слов в документе"""
    substring_map = {
        "WORK_TEXT_POSITION": [],
        "DDS_TEXT_POSITION": [],
        "DATE_TEXT_POSITION": []
    }
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if WORK_NAME_SUBSTRING in text:
            substring_map["WORK_TEXT_POSITION"].append(i)
        if DDS_SUBSTRING in text:
            substring_map["DDS_TEXT_POSITION"].append(i)
        if DATE_SUBSTRING in text:
            substring_map["DATE_TEXT_POSITION"].append(i)
            
    return substring_map

def generated_wfile(doc, substring_map, work_name, dds, date_str=None):
    """Генерирует Word-файл с заменой ключевых слов"""
    # Создаем копию документа в памяти
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    new_doc = Document(file_stream)
    
    # Устанавливаем текущую дату по умолчанию
    if date_str is None:
        date_str = datetime.now().strftime("%d.%m.%Y")
    
    # Выполняем замены
    for key, para_indices in substring_map.items():
        for idx in para_indices:
            if idx < len(new_doc.paragraphs):
                paragraph = new_doc.paragraphs[idx]
                text = paragraph.text
                
                if key == "WORK_TEXT_POSITION" and WORK_NAME_SUBSTRING in text:
                    paragraph.text = text.replace(WORK_NAME_SUBSTRING, work_name)
                
                elif key == "DDS_TEXT_POSITION" and DDS_SUBSTRING in text:
                    paragraph.text = text.replace(DDS_SUBSTRING, dds)
                
                elif key == "DATE_TEXT_POSITION" and DATE_SUBSTRING in text:
                    paragraph.text = text.replace(DATE_SUBSTRING, date_str)
    
    # Сохраняем документ
    out_dir = "out"
    if not os.path.exists(out_dir):
        os.makedirs(out_dir)
    
    filename = f"request_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(out_dir, filename)
    new_doc.save(filepath)
    
    return filepath